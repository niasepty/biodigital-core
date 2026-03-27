import streamlit as st
import pandas as pd
import plotly.express as px
from Bio.Seq import Seq
from Bio.SeqUtils import gc_fraction
from Bio.SeqUtils.ProtParam import ProteinAnalysis
import io
from docx import Document
import streamlit.components.v1 as components
import random

# ==========================================
# BANK QUOTES AESTHETIC BIOLOGI
# ==========================================
KUMPULAN_QUOTES = [
    "\"Mengungkap rahasia kehidupan, satu baris data pada satu waktu.\" 🌿🔬",
    "\"Di antara jutaan basa nukleotida, tersimpan simfoni kehidupan yang menunggu untuk dibaca.\" 🧬✨",
    "\"Setiap sel memiliki ceritanya sendiri. Tugas kita hanyalah mendengarkan.\" 🦠🌌",
    "\"Menerjemahkan bahasa alam semesta, dari untaian DNA hingga algoritma komputasi.\" 🌍📊",
    "\"Keajaiban terbesar tak selalu kasatmata; terkadang ia bersembunyi di balik lensa mikroskop.\" 🔍✨"
]

# ==========================================
# 0. INISIALISASI MEMORI SISTEM (SESSION STATE)
# ==========================================
if 'logged_in' not in st.session_state:
    st.session_state['logged_in'] = False
if 'play_audio' not in st.session_state:
    st.session_state['play_audio'] = False
if 'memori_nutrisi_excel' not in st.session_state:
    st.session_state['memori_nutrisi_excel'] = []
if 'active_user' not in st.session_state:
    st.session_state['active_user'] = ""

# ANTI-BUG QUOTES: Kalau quote-nya kosong, langsung kocok secara otomatis!
if 'current_quote' not in st.session_state or st.session_state['current_quote'] == "":
    st.session_state['current_quote'] = random.choice(KUMPULAN_QUOTES)

# ==========================================
# DATABASE USER (AKUN LOGIN)
# ==========================================
DATABASE_AKUN = {
    "admin": "rahasia123",
    "dosen": "biologi100",
    "asisten": "lab2026",
    "rizky": "tugas123",
    "mahasiswa": "tugas123",
    "nia": "nia123"
}

# ==========================================
# 1. KONFIGURASI HALAMAN & TEMA ENTERPRISE
# ==========================================
st.set_page_config(page_title="BioDigital Core", layout="wide", page_icon="🧬")

st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Plus+Jakarta+Sans:wght@400;600;800;italic&display=swap');
    .stApp { background-color: #0b0f19; color: #e2e8f0; font-family: 'Plus Jakarta Sans', sans-serif; }
    
    .login-box {
        background: linear-gradient(145deg, #111827, #1f2937);
        padding: 40px;
        border-radius: 20px;
        border: 2px solid #3b82f6;
        animation: glowing 3s infinite;
    }
    @keyframes glowing {
        0% { box-shadow: 0 0 10px rgba(59, 130, 246, 0.5); }
        50% { box-shadow: 0 0 30px rgba(16, 185, 129, 0.8); border-color: #10b981; }
        100% { box-shadow: 0 0 10px rgba(59, 130, 246, 0.5); border-color: #3b82f6; }
    }
    
    [data-testid="stMetric"] { background: linear-gradient(145deg, #111827, #1f2937); border: 1px solid #374151; border-radius: 12px; padding: 20px; }
    .title-glow { font-weight: 800; background: linear-gradient(to right, #3b82f6, #10b981); -webkit-background-clip: text; -webkit-text-fill-color: transparent; font-size: 2.8rem; padding-bottom: 5px; margin-bottom: 0;}
    .subtitle-aesthetic { color: #9ca3af; font-size: 1.1rem; font-style: italic; font-weight: 400; margin-top: -10px; margin-bottom: 20px;}
    
    .creator-badge { text-align: center; margin-top: 20px; font-size: 13px; color: #6b7280; }
    .creator-name { color: #10b981; font-weight: 800; font-size: 15px; letter-spacing: 1px; }
    </style>
    """, unsafe_allow_html=True)

# ==========================================
# 2. SISTEM LOGIN (MULTI-USER SECURITY GATE)
# ==========================================
if not st.session_state['logged_in']:
    col1, col2, col3 = st.columns([1, 1.2, 1])
    
    with col2:
        st.markdown("<br><br><br>", unsafe_allow_html=True)
        st.markdown("<div class='login-box'>", unsafe_allow_html=True)
        st.markdown("<h1 style='text-align:center; color:#3b82f6;'>🧬 BioDigital Core</h1>", unsafe_allow_html=True)
        st.markdown("<p style='text-align:center; color:#9ca3af; margin-bottom:30px;'>Silakan Login Sesuai Hak Akses Anda</p>", unsafe_allow_html=True)
        
        user_input = st.text_input("Username", placeholder="Ketik username...")
        pwd_input = st.text_input("Password", type="password", placeholder="Ketik password...")
        
        st.write("") 
        if st.button("🔐 ACCESS SYSTEM", use_container_width=True):
            user_input = user_input.lower().strip() 
            if user_input in DATABASE_AKUN and DATABASE_AKUN[user_input] == pwd_input:
                st.session_state['logged_in'] = True
                st.session_state['play_audio'] = True 
                st.session_state['active_user'] = user_input.capitalize() 
                
                # Mengocok quotes baru saat berhasil login
                st.session_state['current_quote'] = random.choice(KUMPULAN_QUOTES)
                st.rerun() 
            else:
                st.error("❌ AKSES DITOLAK! Username tidak terdaftar atau Password salah.")
        
        st.markdown("</div>", unsafe_allow_html=True)
        
        
        st.markdown("""
            <div class='creator-badge'>
                Developed & Engineered by<br>
                <span class='creator-name'>Nia Septi Tri Pertiwi</span>
            </div>
        """, unsafe_allow_html=True)

# ==========================================
# 3. MAIN DASHBOARD
# ==========================================
else:
    if st.session_state['play_audio']:
        st.balloons() 
        nama_pengguna = st.session_state['active_user']
        
        js_audio = f"""
        <script>
            function bicara() {{
                window.speechSynthesis.cancel();
                setTimeout(function() {{
                    var msg = new SpeechSynthesisUtterance('Akses diterima. Selamat datang {nama_pengguna}. Sistem Bio digital Core siap digunakan.');
                    msg.lang = 'id-ID';
                    msg.rate = 1.0;  
                    msg.pitch = 1.0; 
                    
                    var voices = window.speechSynthesis.getVoices();
                    var indoVoice = voices.find(voice => voice.lang === 'id-ID' || voice.name.includes('Indonesia') || voice.name.includes('Indonesian'));
                    if(indoVoice) msg.voice = indoVoice;
                    
                    window.speechSynthesis.speak(msg);
                }}, 500); 
            }}
            if (speechSynthesis.getVoices().length !== 0) bicara();
            else speechSynthesis.addEventListener('voiceschanged', bicara);
        </script>
        """
        components.html(js_audio, height=0, width=0)
        st.session_state['play_audio'] = False 

    with st.sidebar:
        st.markdown("<h2 style='color: #3b82f6; font-weight:800;'>🧬 BIO-CORE v4.0</h2>", unsafe_allow_html=True)
        st.success(f"👤 Login sebagai: **{st.session_state['active_user']}**")
        st.divider()
        
        menu = st.radio("PILIH MODUL ANALISIS:", ["📈 Kinetika Pertumbuhan", "🧬 Sekuensing Genomik", "🧪 Kalkulator Molaritas"])
        
        st.divider()
        st.markdown("### 🎧 Frekuensi Fokus (BGM)")
        st.markdown("<p style='font-size: 13px; color: #9ca3af; font-style: italic;'>Biarkan harmoni menemani setiap observasi mikroskopismu hari ini. ✨</p>", unsafe_allow_html=True)
        
        playlist = {
            "👽 Sci-Fi Ambient (Eksplorasi Gen)": "https://www.soundhelix.com/examples/mp3/SoundHelix-Song-9.mp3",
            "🤖 Cyberpunk Synth (Analisis Kinetika)": "https://www.soundhelix.com/examples/mp3/SoundHelix-Song-1.mp3",
            "☕ Lofi Chillout (Fase Stasioner)": "https://www.soundhelix.com/examples/mp3/SoundHelix-Song-3.mp3",
            "🚀 Space Groove (Fase Eksponensial)": "https://www.soundhelix.com/examples/mp3/SoundHelix-Song-13.mp3"
        }
        
        opsi_musik = ["🎶 Putar Simfoni Milikmu Sendiri..."] + list(playlist.keys())
        pilihan_lagu = st.selectbox("Pilih Gelombang Suaramu:", opsi_musik)
        
        if pilihan_lagu == "🎶 Putar Simfoni Milikmu Sendiri...":
            lagu_upload = st.file_uploader("Upload Nada Favoritmu (MP3/WAV)", type=['mp3', 'wav'])
            if lagu_upload is not None:
                st.audio(lagu_upload, format="audio/mp3")
                st.success("✨ Nada berhasil diselaraskan! Silakan klik Play.")
            else:
                st.info("Menunggu alunan nadamu...")
        else:
            st.audio(playlist[pilihan_lagu], format="audio/mp3")
        
        st.divider()
        if st.button("🚪 LOGOUT SYSTEM", use_container_width=True):
            st.session_state['logged_in'] = False
            st.session_state['memori_nutrisi_excel'] = []
            st.session_state['active_user'] = ""
            st.rerun()
            
        st.markdown("""
            <div style='margin-top: 40px; padding: 15px; background: rgba(16, 185, 129, 0.05); border-radius: 10px; border-left: 3px solid #10b981;'>
                <span style='font-size: 11px; color: #9ca3af;'>SYSTEM ARCHITECT:</span><br>
                <b style='color: #10b981; font-size: 14px;'>Nia Septi Tri Pertiwi</b>
            </div>
        """, unsafe_allow_html=True)

    # ==========================================
    # QUOTES SELALU MUNCUL DI SINI!
    # ==========================================
    st.markdown("<h1 class='title-glow'>BioDigital Intelligence System</h1>", unsafe_allow_html=True)
    st.markdown(f"<p class='subtitle-aesthetic'>{st.session_state['current_quote']}</p>", unsafe_allow_html=True)
    st.markdown("---")

    # --- MODUL 1: KINETIKA PERTUMBUHAN ---
    if menu == "📈 Kinetika Pertumbuhan":
        st.subheader("📊 Analisis Populasi & Kinetika Mikroba")
        file = st.file_uploader("Unggah Dataset Lab (.xlsx)", type=["xlsx"])

        if file:
            try:
                df = pd.read_excel(file)
                
                nutrisi_unik = df['Jenis_Nutrisi'].dropna().unique().tolist()
                st.session_state['memori_nutrisi_excel'] = [n for n in nutrisi_unik if str(n).lower() != "tanpa_nutrisi"]
                
                q1 = df['Jumlah_Bakteri'].quantile(0.25)
                q3 = df['Jumlah_Bakteri'].quantile(0.75)
                rata_rata = df['Jumlah_Bakteri'].mean()
                
                def kalibrasi_status(val):
                    if val >= q3: return "🟢 Pertumbuhan Cepat (Optimal)"
                    elif val >= q1: return "🟡 Pertumbuhan Normal (Stabil)"
                    else: return "🔴 Pertumbuhan Lambat (Kritis)"
                    
                df['Keterangan'] = df['Jumlah_Bakteri'].apply(kalibrasi_status)

                col_m1, col_m2, col_m3, col_m4 = st.columns(4)
                col_m1.metric("Total Sampel", f"{len(df)} Data")
                col_m2.metric("Rata-rata Populasi", f"{rata_rata:,.2f}")
                col_m3.metric("Batas Optimal (Q3)", f"> {q3:,.2f}")
                col_m4.metric("Batas Kritis (Q1)", f"< {q1:,.2f}")

                st.write("") 

                col_graph, col_table = st.columns([1.8, 1.2])

                with col_graph:
                    st.markdown("#### 📊 Rata-rata Pertumbuhan (Diagram Batang)")
                    df_avg = df.groupby(['Waktu_Jam', 'Jenis_Nutrisi'])['Jumlah_Bakteri'].mean().reset_index()
                    fig = px.bar(df_avg, x='Waktu_Jam', y='Jumlah_Bakteri', color='Jenis_Nutrisi', barmode='group', text_auto='.0f', color_discrete_sequence=px.colors.qualitative.Pastel)
                    fig.update_layout(plot_bgcolor='rgba(0,0,0,0)', paper_bgcolor='rgba(0,0,0,0)', font=dict(color="#9ca3af"), xaxis=dict(showgrid=False, title="Waktu (Jam)", type='category'), yaxis=dict(gridcolor='#374151', title="Rata-rata Sel / ml", zeroline=False), legend=dict(title="Nutrisi", orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1), margin=dict(t=10, l=0, r=0, b=0), height=450)
                    st.plotly_chart(fig, use_container_width=True)

                with col_table:
                    st.markdown("#### 📑 Ledger Data Tersaring")
                    def style_keterangan(val):
                        if 'Optimal' in str(val): return 'color: #10b981; font-weight: bold;'
                        elif 'Kritis' in str(val): return 'color: #ef4444; font-weight: bold;'
                        return 'color: #f59e0b;'
                    styled_df = df.style.format({"Jumlah_Bakteri": "{:,.0f}"}).map(style_keterangan, subset=['Keterangan'])
                    st.dataframe(styled_df, use_container_width=True, hide_index=True)

                st.markdown("---")
                st.markdown("### 🧬 Laporan Profiling Substrat (Auto-Generated)")
                
                ranking_nutrisi = df.groupby('Jenis_Nutrisi')['Jumlah_Bakteri'].mean().sort_values(ascending=False)
                peringkat = 1
                total_nutrisi = len(ranking_nutrisi)
                
                for nutrisi, avg_val in ranking_nutrisi.items():
                    if peringkat == 1:
                        st.success(f"**🏆 Peringkat {peringkat}: {nutrisi} (Fase Eksponensial Optimal)**\n\nMenghasilkan rata-rata pertumbuhan tertinggi (**{avg_val:,.0f} sel/ml**). Substrat ini sangat efisien dimetabolisme oleh bakteri.")
                    elif peringkat == total_nutrisi:
                        st.error(f"**⚠️ Peringkat {peringkat}: {nutrisi} (Kritis / Fase Stasioner)**\n\nMenghasilkan rata-rata pertumbuhan terendah (**{avg_val:,.0f} sel/ml**). Bakteri sangat kesulitan memecah substrat ini.")
                    else:
                        st.info(f"**📊 Peringkat {peringkat}: {nutrisi} (Jalur Energi Alternatif)**\n\nMenghasilkan rata-rata pertumbuhan moderat (**{avg_val:,.0f} sel/ml**). Substrat ini tidak seefektif Peringkat 1.")
                    peringkat += 1

                st.markdown("---")
                st.markdown("### 📥 Export Laporan Praktikum (Microsoft Word)")
                
                doc = Document()
                doc.add_heading('Laporan Analisis Kinetika Pertumbuhan Mikroba', 0)
                doc.add_paragraph('Dokumen ini di-generate secara otomatis oleh BioDigital Core v4.0 System.')
                doc.add_heading('1. Ringkasan Statistik Data', level=1)
                doc.add_paragraph(f'• Total Sampel: {len(df)} Data\n• Rata-rata Populasi Keseluruhan: {rata_rata:,.2f} sel/ml')
                
                buffer = io.BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                
                st.download_button(label="📄 Klik di sini untuk Download Laporan (.docx)", data=buffer, file_name="Laporan_Kinetika_BioDigital.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)

            except KeyError:
                st.error("❌ Pastikan file Excel memiliki kolom yang bernama persis: 'Waktu_Jam', 'Jumlah_Bakteri', dan 'Jenis_Nutrisi'.")
        else:
            st.info("💡 Menunggu tetesan data pertamamu... Silakan unggah file Excel praktikum untuk memulai keajaiban komputasi.")

    # --- MODUL 2: ANALISIS GENOMIK ---
    elif menu == "🧬 Sekuensing Genomik":
        st.subheader("🧬 Analisis Central Dogma & Profiling Asam Amino")
        sumber_dna = st.radio("Sumber Sekuens DNA:", ["Input Manual", "Ekstraksi Virtual dari Data Excel (Smart Match)"], horizontal=True)
        
        dna_sequence = "ATGCGTACGTAGCTAGCTAGCTAGCTAG" 
        nama_gen = "Sekuens Manual"

        if sumber_dna == "Ekstraksi Virtual dari Data Excel (Smart Match)":
            if len(st.session_state['memori_nutrisi_excel']) > 0:
                pilih_nutrisi = st.selectbox("Pilih Nutrisi dari Data Excel Anda:", st.session_state['memori_nutrisi_excel'])
                db_gen = {
                    "Glukosa": ("Gen Glukokinase (glk)", "ATGACTAAACTCCCCCTTGGCGGTGGACCCAGCTCGGCG"),
                    "Laktosa": ("Gen Beta-galaktosidase (lacZ)", "ATGACCATGATTACGGATTCACTGGCCGTCGTTTTACAA")
                }
                kunci_nutrisi = next((key for key in db_gen.keys() if key.lower() in str(pilih_nutrisi).lower()), None)
                if kunci_nutrisi:
                    nama_gen, dna_sequence = db_gen[kunci_nutrisi]
                    st.success(f"✔️ Terdeteksi! Menarik sekuens **{nama_gen}**.")
            else:
                st.warning("⚠️ Belum ada data Excel di Modul 1.")
                sumber_dna = "Input Manual"

        st.markdown(f"**Target Analisis:** {nama_gen}")
        dna_input = st.text_area("Sekuens Basa Nitrogen:", value=dna_sequence, height=100)
        
        if st.button("JALANKAN ALGORITMA MOLEKULER", use_container_width=True):
            seq = Seq(dna_input.upper().strip())
            if not set(str(seq)).issubset({"A", "T", "G", "C", "N"}):
                st.error("❌ Input tidak valid!")
            else:
                prot = seq.translate(to_stop=False)
                clean_prot = str(prot).replace("*", "")
                analysed_p = ProteinAnalysis(clean_prot)
                
                col1, col2 = st.columns(2)
                with col1:
                    st.info(f"**Panjang Untai:** {len(seq)} bp")
                    st.text_area("Hasil mRNA:", str(seq.transcribe()), height=100, disabled=True)
                with col2:
                    if len(clean_prot) > 0:
                        st.success(f"**Massa Molekul:** {analysed_p.molecular_weight():.2f} Dalton")
                    st.text_area("Translasi Protein:", str(prot), height=100, disabled=True)

    # --- MODUL 3: KALKULATOR MOLARITAS ---
    elif menu == "🧪 Kalkulator Molaritas":
        st.subheader("🧪 Presisi Persiapan Reagen Lab")
        st.markdown("Menggunakan standar rumus: m = (M × V × Mr) / 1000")
        
        with st.container():
            st.markdown("<div style='background: #111827; padding: 40px; border-radius:15px; border: 1px solid #374151;'>", unsafe_allow_html=True)
            
            senyawa_db = {"Garam Dapur (NaCl)": 58.44, "Sodium Hidroksida (NaOH)": 40.00, "Glukosa (C6H12O6)": 180.16}
            pilihan_dropdown = []
            
            if len(st.session_state['memori_nutrisi_excel']) > 0:
                for n in st.session_state['memori_nutrisi_excel']: pilihan_dropdown.append(f"🧬 Dari Excel: {n}")
            
            pilihan_dropdown.extend(list(senyawa_db.keys()))
            pilihan_dropdown.append("➕ Input Manual (Zat Lain)")
            pilihan = st.selectbox("1. Pilih Senyawa Terlarut:", pilihan_dropdown)
            
            mr = 100.0 
            if pilihan == "➕ Input Manual (Zat Lain)":
                mr = st.number_input("Masukkan Berat Molekul (Mr) secara manual:", value=1.0)
            elif pilihan in senyawa_db:
                mr = senyawa_db[pilihan]
                
            st.markdown("---")
            col_m, col_v = st.columns(2)
            M = col_m.number_input("2. Target Molaritas / Konsentrasi (M):", value=0.5, step=0.1)
            V = col_v.number_input("3. Target Volume Akhir (ml):", value=250.0, step=10.0)
            
            massa_gram = (M * V * mr) / 1000
            
            st.markdown(f"""
                <div style="text-align: center; margin-top: 30px; padding: 20px; background: rgba(16, 185, 129, 0.1); border-radius: 10px; border: 1px dashed #10b981;">
                    <h4 style="color: #9ca3af; margin-bottom: 5px;">Massa Serbuk yang Harus Ditimbang:</h4>
                    <h1 style="color: #10b981; font-size: 3.5rem; font-weight: 800; margin: 0;">{massa_gram:.4f} Gram</h1>
                </div>
            """, unsafe_allow_html=True)
            st.markdown("</div>", unsafe_allow_html=True)